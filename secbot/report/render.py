"""Report render helpers: Markdown (canonical), HTML, DOCX, PDF.

Spec: `.trellis/spec/backend/report-pipeline.md` §3.

Markdown is the canonical format. HTML is derived from Markdown for PDF
rendering via WeasyPrint. DOCX is built programmatically from the
:class:`ReportModel` because python-docx has no native Markdown ingest.
"""

from __future__ import annotations

import html as _html_mod
from pathlib import Path
from typing import TYPE_CHECKING

from secbot.report.builder import SEVERITY_ORDER, ReportModel, ReportRenderError


def _esc(s: str) -> str:
    """HTML-escape a string (short alias to keep render lines readable)."""
    return _html_mod.escape(str(s), quote=True)


if TYPE_CHECKING:  # pragma: no cover
    pass


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


_SEV_LABELS: dict[str, str] = {
    "critical": "严重",
    "high": "高危",
    "medium": "中危",
    "low": "低危",
    "info": "信息",
}


def _fmt_dt(dt) -> str:
    return dt.isoformat() if dt else "—"


def render_markdown(model: ReportModel) -> str:
    """Render *model* to Markdown (canonical)."""
    out: list[str] = []
    out.append("# 安全扫描报告")
    out.append("")
    out.append(f"- **扫描 ID**: `{model.scan_id}`")
    out.append(f"- **目标**: {model.target}")
    out.append(f"- **开始时间**: {_fmt_dt(model.started_at)}")
    out.append(f"- **结束时间**: {_fmt_dt(model.finished_at)}")
    out.append("")
    out.append("## 摘要")
    out.append("")
    out.append(f"- 资产: **{model.summary.asset_count}**")
    out.append(f"- 服务: **{model.summary.service_count}**")
    out.append(f"- 发现: **{model.summary.finding_count}**")
    out.append("")
    out.append("| 严重级别 | 数量 |")
    out.append("|---|---|")
    for sev in SEVERITY_ORDER:
        out.append(f"| {_SEV_LABELS.get(sev, sev)} | {model.summary.severity_counts.get(sev, 0)} |")
    out.append("")

    if model.is_empty():
        out.append("_本次扫描未记录任何资产。_")
        out.append("")
    else:
        out.append("## 资产")
        out.append("")
        for a in model.assets:
            label = a.hostname or a.ip or a.target
            out.append(f"### {label}")
            out.append("")
            out.append(f"- 目标: `{a.target}`")
            if a.ip:
                out.append(f"- IP: `{a.ip}`")
            if a.hostname:
                out.append(f"- 主机名: {a.hostname}")
            if a.os_guess:
                out.append(f"- 操作系统推测: {a.os_guess}")
            out.append("")

            if a.services:
                out.append("#### 开放服务")
                out.append("")
                out.append("| 端口 | 协议 | 服务 | 产品 | 版本 |")
                out.append("|---|---|---|---|---|")
                for s in a.services:
                    out.append(
                        f"| {s.port} | {s.protocol} | {s.service or '—'} | "
                        f"{s.product or '—'} | {s.version or '—'} |"
                    )
                out.append("")

            if a.findings:
                out.append("#### 漏洞详情")
                out.append("")
                for f in a.findings:
                    sev_label = _SEV_LABELS.get(f.severity, f.severity)
                    out.append(f"##### [{sev_label}] {f.title}")
                    out.append("")
                    meta = f"- **类别**: {f.category}  "
                    meta += f"**发现工具**: {f.discovered_by}"
                    if f.cve_id:
                        meta += f"  \n- **CVE**: `{f.cve_id}`"
                    out.append(meta)
                    if f.affected_url:
                        out.append(f"- **受影响端点**: `{f.affected_url}`")
                    if f.evidence_summary:
                        out.append(f"- **漏洞描述**: {f.evidence_summary}")
                    out.append("")
                    if f.verification_steps:
                        out.append("**验证步骤**:")
                        out.append("")
                        for i, step in enumerate(f.verification_steps, 1):
                            out.append(f"{i}. {step}")
                        out.append("")
                    if f.evidence_detail:
                        out.append("**PoC 证据详情**:")
                        out.append("")
                        out.append("```")
                        out.append(f.evidence_detail)
                        out.append("```")
                        out.append("")
                    if f.remediation:
                        out.append(f"**修复建议**: {f.remediation}")
                        out.append("")
                    if f.references:
                        out.append("**参考资料**:")
                        out.append("")
                        for ref in f.references:
                            out.append(f"- {ref}")
                        out.append("")
                    out.append("---")
                    out.append("")

    if model.appendix.raw_log_paths:
        out.append("## 附录：原始日志")
        out.append("")
        for p in model.appendix.raw_log_paths:
            out.append(f"- `{p}`")
        out.append("")

    return "\n".join(out) + "\n"


# ---------------------------------------------------------------------------
# HTML (for PDF)
# ---------------------------------------------------------------------------


def render_html(model: ReportModel) -> str:
    """Render the canonical HTML report.

    Adopts the ``123.html`` design language: a neutral slate palette, card
    based layout, mini-labels and rounded badges. All styling is inlined
    (no CDN, no JS dependency beyond the print button) so the file opens
    offline and renders identically through WeasyPrint PDF conversion.
    """
    css = """
    * { box-sizing: border-box; }
    body {
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, "PingFang SC", "Microsoft YaHei", sans-serif;
      color: #0f172a;
      background: #fafaf9;
      margin: 0;
      padding: 0;
      line-height: 1.6;
    }
    .mini-label { font-size: 10px; text-transform: uppercase; letter-spacing: .08em; color: #64748b; }
    .seam-line { border-top: 2px dashed #94a3b8; }
    .toolbar {
      position: sticky;
      top: 0;
      z-index: 100;
      display: flex;
      align-items: center;
      justify-content: space-between;
      padding: 12px 24px;
      background: linear-gradient(135deg, #0f172a, #1e293b);
      border-bottom: 1px solid #334155;
    }
    .toolbar-title { color: #f8fafc; font-size: 14px; font-weight: 600; letter-spacing: 0.3px; }
    .btn {
      display: inline-flex; align-items: center; gap: 6px;
      padding: 8px 16px; border-radius: 8px;
      border: 1px solid #475569; background: #1e293b; color: #e2e8f0;
      font-size: 13px; font-weight: 500; cursor: pointer;
      transition: all 0.2s ease; text-decoration: none;
    }
    .btn:hover { background: #334155; border-color: #64748b; }
    .btn svg { width: 14px; height: 14px; fill: currentColor; }
    .container { max-width: 1024px; margin: 0 auto; padding: 40px 24px; }
    .report-header { margin-bottom: 28px; }
    .report-header h1 { font-size: 34px; font-weight: 600; letter-spacing: -0.5px; margin: 6px 0 4px; color: #0f172a; }
    .report-header .subtitle { color: #475569; font-size: 14px; margin: 0; }
    .meta-grid { display: grid; gap: 10px; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); margin-top: 18px; }
    .meta-item { border: 1px solid #e2e8f0; background: #fff; border-radius: 10px; padding: 10px 14px; }
    .meta-item .meta-value { font-size: 13px; color: #0f172a; margin-top: 2px; word-break: break-all; }
    .kpi-grid { display: grid; gap: 14px; grid-template-columns: repeat(auto-fit, minmax(160px, 1fr)); margin-bottom: 22px; }
    .kpi-card { border: 1px solid #e2e8f0; background: #fff; border-radius: 12px; padding: 18px 18px 16px; box-shadow: 0 1px 3px rgba(0,0,0,0.04); }
    .kpi-card .kpi-num { font-size: 30px; font-weight: 600; line-height: 1; display: block; margin-top: 8px; color: #0f172a; }
    .kpi-card.danger { border-color: #ef4444; background: #fef2f2; }
    .kpi-card.danger .kpi-num { color: #b91c1c; }
    .card { border: 1px solid #e2e8f0; background: #fff; border-radius: 14px; padding: 22px 24px; box-shadow: 0 1px 3px rgba(0,0,0,0.04); margin-bottom: 22px; }
    .card > h2 { font-size: 20px; font-weight: 600; margin: 6px 0 16px; color: #0f172a; }
    .sev-bar { display: flex; width: 100%; height: 14px; border-radius: 999px; overflow: hidden; background: #f1f5f9; }
    .sev-bar > span { display: block; height: 100%; }
    .sev-legend { display: flex; flex-wrap: wrap; gap: 14px 22px; margin-top: 16px; }
    .sev-legend .item { display: flex; align-items: center; gap: 8px; font-size: 13px; color: #334155; }
    .sev-legend .dot { width: 10px; height: 10px; border-radius: 3px; display: inline-block; }
    .sev-legend .cnt { font-weight: 600; color: #0f172a; }
    .empty-pill { display: inline-block; padding: 6px 14px; border-radius: 999px; background: #f1f5f9; color: #64748b; font-size: 13px; }
    .asset-card .asset-head { display: flex; flex-wrap: wrap; align-items: baseline; gap: 10px; margin-bottom: 4px; }
    .asset-card .asset-name { font-size: 18px; font-weight: 600; color: #0f172a; }
    .asset-card .asset-sub { font-size: 12px; color: #64748b; }
    .asset-card .asset-meta { display: flex; flex-wrap: wrap; gap: 6px 18px; font-size: 13px; color: #475569; margin: 8px 0 18px; }
    .asset-card .asset-meta code { font-size: 12px; }
    .section-label { font-size: 11px; text-transform: uppercase; letter-spacing: .06em; color: #64748b; font-weight: 600; margin: 18px 0 10px; }
    table { border-collapse: collapse; width: 100%; background: #fff; border: 1px solid #e2e8f0; border-radius: 10px; overflow: hidden; }
    th, td { border-bottom: 1px solid #e2e8f0; padding: 9px 14px; font-size: 13px; text-align: left; }
    tr:last-child td { border-bottom: none; }
    th { background: #f8fafc; font-weight: 600; color: #475569; }
    .badge { padding: 3px 11px; border-radius: 999px; color: #fff; font-size: 12px; font-weight: 500; display: inline-block; line-height: 1.5; }
    .sev-critical { background: #991b1b; }
    .sev-high { background: #dc2626; }
    .sev-medium { background: #d97706; }
    .sev-low { background: #2563eb; }
    .sev-info { background: #475569; }
    code { background: #f1f5f9; padding: 2px 6px; border-radius: 5px; font-size: 12px; color: #334155; word-break: break-word; }
    .finding-card { background: #fff; border-radius: 12px; border: 1px solid #e2e8f0; padding: 16px 20px; margin-bottom: 14px; }
    .finding-header { display: flex; align-items: center; gap: 10px; margin-bottom: 8px; }
    .finding-title { font-weight: 600; font-size: 15px; color: #0f172a; }
    .finding-meta { font-size: 12px; color: #64748b; }
    .finding-section { margin-top: 12px; }
    .finding-section-title { font-weight: 600; font-size: 10px; color: #64748b; text-transform: uppercase; letter-spacing: .08em; margin-bottom: 4px; }
    .finding-detail { font-size: 13px; color: #334155; white-space: pre-wrap; word-break: break-word; }
    .finding-steps { padding-left: 18px; margin: 4px 0 0 0; }
    .finding-steps li { font-size: 13px; color: #334155; margin-bottom: 3px; }
    .finding-refs { font-size: 12px; }
    .finding-refs a { color: #2563eb; text-decoration: none; }
    .finding-refs a:hover { text-decoration: underline; }
    pre.evidence { background: #0f172a; color: #e2e8f0; padding: 12px 16px; border-radius: 8px; font-size: 12px; white-space: pre-wrap; word-break: break-all; max-height: 600px; overflow: auto; margin: 0; }
    @media print {
      .toolbar { display: none !important; }
      body { background: #fff; }
      .container { padding: 0; max-width: none; }
      .card, .kpi-card, .finding-card, .asset-card { box-shadow: none; break-inside: avoid; }
    }
    """

    # --- severity aggregation for the visual summary -------------------
    sev_counts = {sev: model.summary.severity_counts.get(sev, 0) for sev in SEVERITY_ORDER}
    sev_total = sum(sev_counts.values())
    sev_color = {
        "critical": "#991b1b",
        "high": "#dc2626",
        "medium": "#d97706",
        "low": "#2563eb",
        "info": "#475569",
    }
    critical_count = sev_counts.get("critical", 0)
    high_count = sev_counts.get("high", 0)

    lines: list[str] = [
        "<!DOCTYPE html>",
        '<html lang="zh-CN"><head><meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        "<title>安全扫描报告</title>",
        f"<style>{css}</style></head><body>",
        '<div class="toolbar">',
        '  <span class="toolbar-title">安全扫描报告</span>',
        '  <button class="btn" onclick="window.print()" title="打印 / 导出 PDF">',
        '    <svg viewBox="0 0 24 24"><path d="M19 8H5c-1.66 0-3 1.34-3 3v6h4v4h12v-4h4v-6c0-1.66-1.34-3-3-3zm-3 11H8v-5h8v5zm3-7c-.55 0-1-.45-1-1s.45-1 1-1 1 .45 1 1-.45 1-1 1zm-1-9H6v4h12V3z"/></svg>',
        '    打印 / 导出 PDF',
        "  </button>",
        "</div>",
        '<main class="container">',
        '  <header class="report-header">',
        '    <p class="mini-label">Security scan report</p>',
        "    <h1>安全扫描报告</h1>",
        f'    <p class="subtitle">目标 {_esc(model.target)}</p>',
        '    <div class="meta-grid">',
        '      <div class="meta-item"><p class="mini-label">扫描 ID</p>'
        f'<div class="meta-value"><code>{_esc(model.scan_id)}</code></div></div>',
        '      <div class="meta-item"><p class="mini-label">目标</p>'
        f'<div class="meta-value">{_esc(model.target)}</div></div>',
        '      <div class="meta-item"><p class="mini-label">开始时间</p>'
        f'<div class="meta-value">{_esc(_fmt_dt(model.started_at))}</div></div>',
        '      <div class="meta-item"><p class="mini-label">结束时间</p>'
        f'<div class="meta-value">{_esc(_fmt_dt(model.finished_at))}</div></div>',
        "    </div>",
        "  </header>",
    ]

    # --- KPI cards ----------------------------------------------------
    lines.append('<section class="kpi-grid">')
    kpis = [
        ("资产", model.summary.asset_count, False),
        ("服务", model.summary.service_count, False),
        ("漏洞发现", model.summary.finding_count, False),
        ("严重 / 高危", critical_count + high_count, (critical_count + high_count) > 0),
    ]
    for label, value, danger in kpis:
        cls = "kpi-card danger" if danger else "kpi-card"
        lines.append(
            f'<div class="{cls}"><p class="mini-label">{label}</p>'
            f'<span class="kpi-num">{value}</span></div>'
        )
    lines.append("</section>")

    # --- severity distribution (pure CSS) -----------------------------
    lines.append('<section class="card">')
    lines.append('<p class="mini-label">风险分布</p>')
    if sev_total > 0:
        lines.append('<div class="sev-bar">')
        for sev in SEVERITY_ORDER:
            cnt = sev_counts[sev]
            if cnt <= 0:
                continue
            pct = cnt / sev_total * 100
            lines.append(
                f'<span style="width:{pct:.2f}%;background:{sev_color[sev]}" '
                f'title="{_SEV_LABELS.get(sev, sev)}: {cnt}"></span>'
            )
        lines.append("</div>")
        lines.append('<div class="sev-legend">')
        for sev in SEVERITY_ORDER:
            cnt = sev_counts[sev]
            lines.append(
                '<span class="item">'
                f'<span class="dot" style="background:{sev_color[sev]}"></span>'
                f'{_SEV_LABELS.get(sev, sev)} <span class="cnt">{cnt}</span></span>'
            )
        lines.append("</div>")
    else:
        lines.append('<div class="empty-pill">本次扫描未发现漏洞</div>')
    lines.append("</section>")

    # --- per-asset detail cards ---------------------------------------
    for a in model.assets:
        label = a.hostname or a.ip or a.target
        lines.append('<article class="card asset-card">')
        lines.append('<div class="asset-head">')
        lines.append(f'<span class="asset-name">{_esc(label)}</span>')
        if a.findings:
            lines.append(f'<span class="asset-sub">{len(a.findings)} 项发现</span>')
        lines.append("</div>")
        lines.append('<div class="asset-meta">')
        lines.append(f'<span>目标 <code>{_esc(a.target)}</code></span>')
        if a.ip:
            lines.append(f'<span>IP <code>{_esc(a.ip)}</code></span>')
        if a.os_guess:
            lines.append(f'<span>操作系统 {_esc(a.os_guess)}</span>')
        lines.append("</div>")
        if a.services:
            lines.append('<p class="section-label">开放服务</p>')
            lines.append("<table><thead><tr><th>端口</th><th>协议</th>"
                         "<th>服务</th><th>产品</th><th>版本</th></tr></thead><tbody>")
            for s in a.services:
                lines.append(
                    f"<tr><td>{s.port}</td><td>{_esc(s.protocol)}</td>"
                    f"<td>{_esc(s.service) if s.service else '—'}</td>"
                    f"<td>{_esc(s.product) if s.product else '—'}</td>"
                    f"<td>{_esc(s.version) if s.version else '—'}</td></tr>"
                )
            lines.append("</tbody></table>")
        if a.findings:
            lines.append('<p class="section-label">漏洞详情</p>')
            for f in a.findings:
                sev_label = _SEV_LABELS.get(f.severity, f.severity)
                lines.append('<div class="finding-card">')
                # Header
                lines.append('<div class="finding-header">')
                lines.append(f'<span class="badge sev-{f.severity}">{sev_label}</span>')
                lines.append(f'<span class="finding-title">{_esc(f.title)}</span>')
                lines.append('</div>')
                # Meta row
                meta_parts = [f'类别: {f.category}', f'发现工具: {f.discovered_by}']
                if f.cve_id:
                    meta_parts.append(f'CVE: {f.cve_id}')
                lines.append(f'<div class="finding-meta">{" · ".join(_esc(p) for p in meta_parts)}</div>')
                # Affected URL
                if f.affected_url:
                    lines.append('<div class="finding-section">')
                    lines.append('<div class="finding-section-title">受影响端点</div>')
                    lines.append(f'<code>{_esc(f.affected_url)}</code>')
                    lines.append('</div>')
                # Evidence summary
                if f.evidence_summary:
                    lines.append('<div class="finding-section">')
                    lines.append('<div class="finding-section-title">漏洞描述</div>')
                    lines.append(f'<div class="finding-detail">{_esc(f.evidence_summary)}</div>')
                    lines.append('</div>')
                # Verification steps
                if f.verification_steps:
                    lines.append('<div class="finding-section">')
                    lines.append('<div class="finding-section-title">验证步骤</div>')
                    lines.append('<ol class="finding-steps">')
                    for step in f.verification_steps:
                        lines.append(f'<li>{_esc(step)}</li>')
                    lines.append('</ol>')
                    lines.append('</div>')
                # Evidence detail (structured PoC: request/response/curl)
                if f.evidence_raw and isinstance(f.evidence_raw, dict):
                    ev = f.evidence_raw
                    # Description
                    ev_desc = ev.get("description") or ev.get("raw")
                    if ev_desc and not f.evidence_summary:
                        lines.append('<div class="finding-section">')
                        lines.append('<div class="finding-section-title">漏洞描述</div>')
                        lines.append(f'<div class="finding-detail">{_esc(ev_desc)}</div>')
                        lines.append('</div>')
                    # Matched location
                    ev_matched = ev.get("matched_at") or ev.get("url")
                    if ev_matched:
                        lines.append('<div class="finding-section">')
                        lines.append('<div class="finding-section-title">匹配位置</div>')
                        lines.append(f'<code>{_esc(ev_matched)}</code>')
                        lines.append('</div>')
                    # Payload
                    ev_payload = ev.get("payload")
                    if ev_payload:
                        lines.append('<div class="finding-section">')
                        lines.append('<div class="finding-section-title">攻击载荷</div>')
                        lines.append(f'<pre class="evidence">{_esc(ev_payload)}</pre>')
                        lines.append('</div>')
                    # PoC Request
                    ev_req = ev.get("request")
                    if ev_req:
                        lines.append('<div class="finding-section">')
                        lines.append('<div class="finding-section-title">PoC 请求</div>')
                        lines.append(f'<pre class="evidence">{_esc(str(ev_req))}</pre>')
                        lines.append('</div>')
                    # System Response
                    ev_resp = ev.get("response")
                    if ev_resp:
                        lines.append('<div class="finding-section">')
                        lines.append('<div class="finding-section-title">系统响应 (证据)</div>')
                        lines.append(f'<pre class="evidence">{_esc(str(ev_resp))}</pre>')
                        lines.append('</div>')
                    # Curl command
                    ev_curl = ev.get("curl_command") or ev.get("curl")
                    if ev_curl:
                        lines.append('<div class="finding-section">')
                        lines.append('<div class="finding-section-title">复现命令</div>')
                        lines.append(f'<pre class="evidence">{_esc(ev_curl)}</pre>')
                        lines.append('</div>')
                elif f.evidence_detail:
                    lines.append('<div class="finding-section">')
                    lines.append('<div class="finding-section-title">证据详情</div>')
                    lines.append(f'<pre class="evidence">{_esc(f.evidence_detail)}</pre>')
                    lines.append('</div>')
                # Remediation
                if f.remediation:
                    lines.append('<div class="finding-section">')
                    lines.append('<div class="finding-section-title">修复建议</div>')
                    lines.append(f'<div class="finding-detail">{_esc(f.remediation)}</div>')
                    lines.append('</div>')
                # References
                if f.references:
                    lines.append('<div class="finding-section">')
                    lines.append('<div class="finding-section-title">参考资料</div>')
                    lines.append('<div class="finding-refs">')
                    for ref in f.references:
                        ref_escaped = _esc(ref)
                        if ref.startswith("http"):
                            lines.append(f'<a href="{ref_escaped}" target="_blank">{ref_escaped}</a><br>')
                        else:
                            lines.append(f'{ref_escaped}<br>')
                    lines.append('</div>')
                    lines.append('</div>')
                lines.append('</div>')
        lines.append("</article>")
    lines.append("</main></body></html>")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# PDF (WeasyPrint)
# ---------------------------------------------------------------------------


def render_pdf(model: ReportModel, out_path: Path) -> Path:
    """Render *model* as PDF via WeasyPrint.

    Raises :class:`ReportRenderError` if the optional ``weasyprint`` dep is
    not installed or its system libraries (cairo/pango) are missing.
    """
    try:
        from weasyprint import HTML  # type: ignore
    except ImportError as exc:
        raise ReportRenderError(
            "weasyprint is not installed; run `pip install weasyprint` "
            "(requires cairo/pango system libraries)"
        ) from exc
    except OSError as exc:  # pragma: no cover - env specific
        raise ReportRenderError(
            f"weasyprint failed to load system libs: {exc}"
        ) from exc

    html = render_html(model)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html).write_pdf(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------


def render_docx(model: ReportModel, out_path: Path) -> Path:
    """Render *model* as DOCX via python-docx."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise ReportRenderError(
            "python-docx is not installed; run `pip install python-docx`"
        ) from exc

    doc = Document()
    doc.add_heading("安全扫描报告", level=1)
    p = doc.add_paragraph()
    p.add_run("扫描 ID: ").bold = True
    p.add_run(model.scan_id)
    p = doc.add_paragraph()
    p.add_run("目标: ").bold = True
    p.add_run(model.target)
    doc.add_paragraph(f"开始时间: {_fmt_dt(model.started_at)}")
    doc.add_paragraph(f"结束时间: {_fmt_dt(model.finished_at)}")

    doc.add_heading("摘要", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light List Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "严重级别"
    hdr[1].text = "数量"
    for sev in SEVERITY_ORDER:
        row = tbl.add_row().cells
        row[0].text = _SEV_LABELS.get(sev, sev)
        row[1].text = str(model.summary.severity_counts.get(sev, 0))

    if model.is_empty():
        doc.add_paragraph("本次扫描未记录任何资产。")
    else:
        doc.add_heading("资产", level=2)
        for a in model.assets:
            doc.add_heading(a.hostname or a.ip or a.target, level=3)
            doc.add_paragraph(f"目标: {a.target}")
            if a.ip:
                doc.add_paragraph(f"IP: {a.ip}")
            if a.os_guess:
                doc.add_paragraph(f"操作系统推测: {a.os_guess}")

            if a.findings:
                doc.add_heading("漏洞详情", level=4)
                for f in a.findings:
                    sev_label = _SEV_LABELS.get(f.severity, f.severity)
                    doc.add_heading(f"[{sev_label}] {f.title}", level=5)
                    p = doc.add_paragraph()
                    p.add_run("类别: ").bold = True
                    p.add_run(f.category)
                    p.add_run("  ·  ")
                    p.add_run("发现工具: ").bold = True
                    p.add_run(f.discovered_by)
                    if f.cve_id:
                        p.add_run("  ·  ")
                        p.add_run("CVE: ").bold = True
                        p.add_run(f.cve_id)
                    if f.affected_url:
                        p2 = doc.add_paragraph()
                        p2.add_run("受影响端点: ").bold = True
                        p2.add_run(f.affected_url)
                    if f.evidence_summary:
                        p3 = doc.add_paragraph()
                        p3.add_run("漏洞描述: ").bold = True
                        p3.add_run(f.evidence_summary)
                    if f.verification_steps:
                        doc.add_paragraph("验证步骤:").bold = True
                        for i, step in enumerate(f.verification_steps, 1):
                            doc.add_paragraph(f"{i}. {step}", style="List Number")
                    if f.evidence_detail:
                        doc.add_paragraph("PoC 证据详情:").bold = True
                        doc.add_paragraph(f.evidence_detail, style="No Spacing")
                    if f.remediation:
                        p4 = doc.add_paragraph()
                        p4.add_run("修复建议: ").bold = True
                        p4.add_run(f.remediation)
                    if f.references:
                        doc.add_paragraph("参考资料:").bold = True
                        for ref in f.references:
                            doc.add_paragraph(ref, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


__all__ = [
    "render_docx",
    "render_html",
    "render_markdown",
    "render_pdf",
]
